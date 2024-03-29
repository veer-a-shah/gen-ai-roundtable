#Importing required libraries
import openai
from pydub import AudioSegment
from pydub.utils import make_chunks
import json
import tempfile
import os
import spacy
from docx import Document

def processing_func(audio_file_location: str,
                     abstract_location: str, 
                     document_name: str, 
                     openai_model: str = "gpt-4-1106-preview"):
    
    #Checking how the audio is being converted
    print("Audio converter being used: " + str(AudioSegment.converter))

    # Set up OpenAI API credentials
    from dotenv import load_dotenv, find_dotenv
    _ = load_dotenv(find_dotenv())
    openai.api_key = os.getenv('OPENAI_API_KEY')

    #Deleting document if it exists
    if os.path.exists(document_name):
        os.remove(document_name)

    #Getting abstract of paper
    with open(abstract_location, 'r', encoding='utf-8') as file:
        abstract = file.read()
    print("\n Abstract: \n" + abstract)

    #Checking if audio file exists
    if os.path.isfile(audio_file_location):
        print("\n Audio file exists and will be opened. \n")
    else:
        print("\n Audio file does not exist. \n")

    #Opening audio file
    audio_file = AudioSegment.from_file(audio_file_location, format="m4a")

    # PyDub handles time in milliseconds
    ten_minutes = 10 * 60 * 1000
    one_minute = 1 * 60 * 1000

    #Just using the first 10 minutes
    #audio_file = audio_file[:one_minute]

    #Dividing audiofile into chunks
    chunks = make_chunks(audio_file, one_minute)

    # Transcribe each chunk and perform diarization
    transcripts = []
    speakers = []

    # Create a temporary directory to store the audio chunks
    temp_dir = tempfile.mkdtemp()

    for i, chunk in enumerate(chunks):
        audio_data = chunk.raw_data

        # Write audio data to temporary file
        with tempfile.NamedTemporaryFile(delete=False, dir=temp_dir, suffix=".wav") as f:
            temp_file_path = f.name
            chunk.export(temp_file_path, format="wav")

        with open(temp_file_path, mode='rb') as f:
            transcript = openai.audio.transcriptions.create(model="whisper-1",file=f,response_format="text")
            transcripts.append(transcript)

        # Remove temporary file
        os.unlink(temp_file_path)

        print("\n Chunk trancription completed: " + str(i) + " of " + str(len(chunks)))

    transcripts = " ".join(transcripts)

    print("\n Transcription: \n" + transcripts)

    def strip_unimportant_parts(text):
        # Load English language model
        nlp = spacy.load("en_core_web_sm")
        # Set of unimportant part-of-speech tags
        unimportant_tags = {"CCONJ", "DET", "PART", "PRON", "SCONJ"}
        # Process the text with spaCy
        doc = nlp(text)
        # Filter out unimportant parts based on part-of-speech tags
        stripped_text = " ".join(token.text for token in doc if token.pos_ not in unimportant_tags)
        return stripped_text

    #Stripping the transcript
    stripped_transcript = strip_unimportant_parts(transcripts)
    print("\n Stripped transcription: \n" + stripped_transcript)

    #Importing and readying prompts
    with open('prompts.json', 'r') as file:
        prompts = json.load(file)

    system_prompt = prompts["system_prompt"]
    summarisation_prompt = prompts["summarisation_prompt"]
    summarisation_prompt = summarisation_prompt.format(abstract = abstract, transcript = stripped_transcript)
    talking_points_prompt = prompts["talking_points_prompt"]
    areas_to_investigate_prompt = prompts["areas_to_investigate_prompt"]

    #Defning LLM processing functions
    def summary_func():
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[{"role": "system", "content": system_prompt},
                {"role": "user", "content": summarisation_prompt}]
        )
        summary = response.choices[0].message.content
        print("\n Summary: \n" + summary)
        return summary


    def talking_points_func(summary):
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[{"role": "system", "content": system_prompt},
                {"role": "user", "content": summarisation_prompt},
                {"role": "assistant", "content": summary},
                {"role": "user", "content": talking_points_prompt}]
        )
        talking_points = response.choices[0].message.content
        print("\n Talking points: \n" + talking_points)
        return talking_points


    def areas_to_investigate_func(summary, talking_points):
        response = openai.chat.completions.create(
            model=openai_model,
            messages=[{"role": "system", "content": system_prompt},
                {"role": "user", "content": summarisation_prompt},
                {"role": "assistant", "content": summary},
                {"role": "user", "content": talking_points_prompt},
                {"role": "assistant", "content": talking_points},
                {"role": "user", "content": areas_to_investigate_prompt}]
        )
        areas_to_investigate = response.choices[0].message.content
        print("\n Alternative areas to investigate: \n" + areas_to_investigate)
        return areas_to_investigate

    #Document creation function
    def write_strings_to_word(summary, talking_points, alternatives, document_name):
        # Create a new Word document
        doc = Document()

        # Add string1 with a title
        doc.add_heading("Summary:", level=2)
        doc.add_paragraph(summary)

        # Add string2 with a title
        doc.add_heading("Talking Points:", level=2)
        doc.add_paragraph(talking_points)

        # Add string3 with a title
        doc.add_heading("Alternative Areas to Investigate:", level=2)
        doc.add_paragraph(alternatives)

        # Save the document
        doc.save(document_name)
        print(f" \n Word document '{document_name}' created successfully.")

    #Running the functions
    summary = summary_func()
    summary.replace("**", "")
    talking_points = talking_points_func(summary)
    talking_points.replace("**", "")
    areas_to_investigate = areas_to_investigate_func(summary, talking_points)
    areas_to_investigate.replace("**", "")

    write_strings_to_word(summary, talking_points, areas_to_investigate, document_name)

sessions = ["Session 1 16-06-2023", 
            "Session 2 04-07-2023",
            "Session 3 25-07-2023",
            "Session 4 08-08-2023",
            "Session 5 16-08-2023",
            "Session 8 13-10-2023"]

for session in sessions:

    audio_file_location = "/Users/veer.a.shah/Code/gen-ai-roundtable/data_files/recordings/{}.m4a".format(session)
    document_name = "/Users/veer.a.shah/Code/gen-ai-roundtable/data_files/session_notes/{}.docx".format(session)
    abstract_location = "/Users/veer.a.shah/Code/gen-ai-roundtable/data_files/abstracts/{} Abstract.txt".format(session)

    print("\n Processing: {} \n".format(session))

    processing_func(audio_file_location=audio_file_location,
                    document_name=document_name,
                    abstract_location=abstract_location,
                    openai_model="gpt-4-1106-preview")
    
print("\n SESSION PROCESSING COMPLETE!")
