import openai
from pydub import AudioSegment
from pydub.utils import make_chunks
import json
import tempfile
import os
import spacy
from docx import Document

#Checking how the audio is being converted
print(AudioSegment.converter)

# Set up OpenAI API credentials
from dotenv import load_dotenv, find_dotenv
_ = load_dotenv(find_dotenv())
openai.api_key = os.getenv('OPENAI_API_KEY')

#Setting audio file location
audio_file_location = "C:/Code/GenAI_Roundtable/Recordings/roundtable-1.m4a"
document_name = "Session 1"

#Getting abstract of paper
with open("abstract.txt", 'r', encoding='utf-8') as file:
    abstract = file.read()
print("Text read from the file:")

#Checking if audio file exists
if os.path.isfile(audio_file_location):
    print("Audio file exists and will be opened.")
else:
    print("Audio file does not exist.")

#Opening audio file
audio_file = AudioSegment.from_file(audio_file_location, format="m4a")

# PyDub handles time in milliseconds
ten_minutes = 10 * 60 * 1000

#Just using the first 10 minutes
audio_file = audio_file[:ten_minutes]

#Dividing audiofile into chunks
chunks = make_chunks(audio_file, ten_minutes)

# Transcribe each chunk and perform diarization
transcripts = []
speakers = []

# Create a temporary directory to store the audio chunks
temp_dir = tempfile.mkdtemp()

for i, chunk in enumerate(chunks):
    audio_data = chunk.raw_data

    # Write audio data to temporary file
    with tempfile.NamedTemporaryFile(delete=False, dir=temp_dir, suffix=".wav") as f:
        temp_file_path = f.name
        chunk.export(temp_file_path, format="wav")

    # Transcribe audio from file
    with open(temp_file_path, mode='rb') as f:
        response = openai.Audio.transcribe("whisper-1", file=f)
        transcript = response['text']
        transcripts.append(transcript)

    # Remove temporary file
    os.unlink(temp_file_path)

    print("Chunk completed: " + str(i))

print(transcripts)

# Create functions to summarise the discussion, convert it into interesting talking points and suggest alternative areas to investigate
#Importing prompts
with open('prompts.json', 'r') as file:
    prompts = json.load(file)

system_prompt = prompts["system_prompt"]
summarisation_prompt = prompts["summarisation_prompt"]
summarisation_prompt = summarisation_prompt.format(abstract = abstract, transcript = transcripts)

print(summarisation_prompt)

def summarise_discussion_test():
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": system_prompt},
            {"role": "user", "content": summarisation_prompt}]
    )
    summary = response["choices"][0]["message"]["content"].strip()
    return summary

# prompts
formal_pov = "Incorporate the valuable insights from the discussion. Identify opportunities to enhance service offerings, optimize operational efficiencies, and capitalize on emerging market trends. Focus on adapting to evolving client needs, leveraging data-driven decision-making, and fostering strategic partnerships. Consider exploring AI-driven solutions and digital transformation to stay competitive in the ever-changing business landscape. You are English."
informal_summary = "Take action based on the fascinating insights discussed. Stay ahead of the curve by adapting to evolving client needs, making data-driven decisions, and embracing digital transformation. Explore strategic partnerships and leverage emerging technologies to gain a competitive edge. Look for opportunities to implement innovative approaches and improve operational efficiency. You are English."
next_steps = "Engage in further exploration based on the thought-provoking discussion. Take action by asking the following questions: How can we effectively adapt to evolving client needs? What are some successful examples of data-driven decision-making in our industry? Identify potential strategic partners to expand our reach and offerings. Consider implementing AI and automation to streamline operations and enhance customer experiences. Stay updated on new market trends, such as sustainability initiatives or customer-centric strategies, and assess their relevance for our business. You are English."

def strip_unimportant_parts(text):
    # Load English language model
    nlp = spacy.load("en_core_web_sm")
    # Set of unimportant part-of-speech tags
    unimportant_tags = {"CCONJ", "DET", "PART", "PRON", "SCONJ"}
    # Process the text with spaCy
    doc = nlp(text)
    # Filter out unimportant parts based on part-of-speech tags
    stripped_text = " ".join(token.text for token in doc if token.pos_ not in unimportant_tags)
    return stripped_text

def summarise_discussion(text):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": formal_pov},
            {"role": "user", "content": text}]
    )
    summary = response["choices"][0]["message"]["content"].strip()
    return summary

def generate_talking_points(text):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": informal_summary},
            {"role": "user", "content": text}]
    )
    talking_points = response["choices"][0]["message"]["content"].strip()
    return talking_points

def suggest_alternative_areas(text):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": informal_summary},
            {"role": "user", "content": text}]
    )
    alternatives = response["choices"][0]["message"]["content"].strip()
    return alternatives

def write_strings_to_word(summary, talking_points, alternatives, document_name):
    # Create a new Word document
    doc = Document()

    # Add string1 with a title
    doc.add_heading("Summary:", level=1)
    doc.add_paragraph(summary)

    # Add string2 with a title
    doc.add_heading("Talking Points:", level=1)
    doc.add_paragraph(talking_points)

    # Add string3 with a title
    doc.add_heading("Alternative Areas to Investigate:", level=1)
    doc.add_paragraph(alternatives)

    # Save the document
    doc.save(document_name)
    print(f"Word document '{document_name}' created successfully.")


text = " ".join(transcripts)
stripped_transcription = strip_unimportant_parts(text)

summary = summarise_discussion(stripped_transcription)
talking_points = generate_talking_points(stripped_transcription)
alternatives = suggest_alternative_areas(stripped_transcription)

write_strings_to_word(summary, talking_points, alternatives, document_name)

print("Summary:")
print(summary)
print("\nTalking Points:")
print(talking_points)
print("\nAlternative Areas to Investigate:")
print(alternatives)